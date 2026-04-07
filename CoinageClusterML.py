# ==========================================================
# IMPERIAL ATOMIC ANALYSIS SUITE 2026 - FINAL CLEAN VERSION (N ≤ 55 ONLY)
# Magic Numbers using Δ²E + ML vs DFT Comparison
# ==========================================================

import os
import re
import time
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import warnings
import shutil                               
from datetime import datetime
from sklearn.model_selection import train_test_split, learning_curve
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import mean_absolute_error, r2_score
from sklearn.ensemble import ExtraTreesRegressor, RandomForestRegressor, GradientBoostingRegressor
from sklearn.neural_network import MLPRegressor
import xgboost as xgb
import lightgbm as lgb
from catboost import CatBoostRegressor
import shap
from scipy.signal import find_peaks
from docx import Document
from docx.shared import Inches

warnings.filterwarnings('ignore')

# =========================================
plt.rcParams.update({
    'font.family': 'serif',
    'font.size': 11,
    'axes.titlesize': 13,
    'figure.figsize': (5.8, 4.3),
    'savefig.dpi': 600
})

ROOT = "IMPERIAL_FINAL_2026_N55"
os.makedirs(f"{ROOT}/Figures", exist_ok=True)
os.makedirs(f"{ROOT}/Tables", exist_ok=True)

print("(N ≤ 55 )...")

# ========================================
df = pd.read_excel("CuAgAu.xlsx")

def detect_metal(text):
    if pd.isna(text): return None
    text = str(text).upper()
    if "AU" in text: return "Au"
    if "AG" in text: return "Ag"
    if "CU" in text: return "Cu"
    return None

df["metal"] = df["structure_xyz"].apply(detect_metal)
metal_map = {"Cu":29, "Ag":47, "Au":79}
df["metal_Z"] = df["metal"].map(metal_map)

def extract_coords(text):
    if pd.isna(text): return None
    pattern = r'(Cu|Ag|Au)\s+([-+]?\d*\.?\d+)\s+([-+]?\d*\.?\d+)\s+([-+]?\d*\.?\d+)'
    matches = re.findall(pattern, str(text))
    if len(matches) < 5: return None
    return np.array([[float(x), float(y), float(z)] for _,x,y,z in matches])

df["coords"] = df["structure_xyz"].apply(extract_coords)
df = df.dropna(subset=["coords"]).reset_index(drop=True)

# ===  N ≤ 55  ===
df = df[df["n_atoms"] <= 55].reset_index(drop=True)
print(f" Final dataset (N ≤ 55): {len(df)} samples")

df["binding_energy_per_atom"] = df["energy_dft"] / df["n_atoms"]

# ========================================
def compute_geometry(coords):
    if len(coords) < 3:
        return [np.nan] * 9
    centroid = np.mean(coords, axis=0)
    dist = np.linalg.norm(coords - centroid, axis=1)
    rg = np.sqrt(np.mean(dist**2))
    asph = (dist.max() - dist.min()) / (dist.mean() + 1e-12)
    return [
        dist.mean(), dist.std(), dist.max(), rg, asph,
        coords[:,0].max() - coords[:,0].min(),
        coords[:,1].max() - coords[:,1].min(),
        coords[:,2].max() - coords[:,2].min(),
        rg / (dist.mean() + 1e-12)
    ]

geo_cols = ["mean_dist", "std_dist", "max_dist", "radius_gyration", "asphericity",
            "bbox_x", "bbox_y", "bbox_z", "compactness"]

print("working...")
geo = np.array([compute_geometry(c) for c in df["coords"]])
for i, col in enumerate(geo_cols):
    df[col] = geo[:, i]

feature_cols = ["metal_Z", "homo_lumo_gap", "n_val_electrons", "magnetic_moment"] + geo_cols

X = df[feature_cols].fillna(df[feature_cols].median())
y = df["binding_energy_per_atom"]

X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.15, random_state=42, stratify=df["metal"])

scaler = StandardScaler()
X_train_s = scaler.fit_transform(X_train)
X_test_s = scaler.transform(X_test)

# ========================================
model_suite = {
    "ExtraTrees": ExtraTreesRegressor(n_estimators=1000, max_depth=15, min_samples_leaf=3, random_state=42, n_jobs=-1),
    "RandomForest": RandomForestRegressor(n_estimators=1000, max_depth=12, min_samples_leaf=4, random_state=42, n_jobs=-1),
    "XGBoost": xgb.XGBRegressor(n_estimators=1000, learning_rate=0.03, max_depth=6, subsample=0.8, colsample_bytree=0.8, reg_alpha=0.1, reg_lambda=1.0, random_state=42),
    "LightGBM": lgb.LGBMRegressor(n_estimators=1000, learning_rate=0.03, max_depth=8, num_leaves=31, random_state=42, verbose=-1),
    "CatBoost": CatBoostRegressor(n_estimators=1000, learning_rate=0.03, depth=7, l2_leaf_reg=5, verbose=0, random_state=42),
    "GradientBoosting": GradientBoostingRegressor(n_estimators=800, learning_rate=0.04, max_depth=5, random_state=42),
    "NeuralNet": MLPRegressor(hidden_layer_sizes=(128, 64), max_iter=1500, alpha=0.05, random_state=42)
}

print("\n=== مقارنة الـ 7 نماذج ===")
results = []
trained_models = {}

for name, model in model_suite.items():
    t0 = time.time()
    model.fit(X_train_s, y_train)
    train_time = time.time() - t0
    pred = model.predict(X_test_s)
    mae = mean_absolute_error(y_test, pred)
    r2 = r2_score(y_test, pred)
    results.append([name, round(mae,5), round(r2,4), round(train_time,2)])
    trained_models[name] = model
    print(f"{name:18} → MAE: {mae:.5f} | R²: {r2:.4f} | Time: {train_time:.2f} sec")

results_df = pd.DataFrame(results, columns=["Model", "MAE", "R2", "Training_Time_sec"])
results_df = results_df.sort_values("MAE")
best_model_name = results_df.iloc[0]["Model"]
best_model = trained_models[best_model_name]

print(f"\n Best_model: {best_model_name} | MAE = {results_df.iloc[0]['MAE']:.5f} | R² = {results_df.iloc[0]['R2']:.4f}")

# =========================================
print("\n=====")
df_test = X_test.copy()
df_test["Actual_DFT"] = y_test.values
df_test["Predicted_ML"] = best_model.predict(X_test_s)
df_test["Error"] = np.abs(df_test["Actual_DFT"] - df_test["Predicted_ML"])
df_test["n_atoms"] = df.loc[X_test.index, "n_atoms"].values
df_test["metal"] = df.loc[X_test.index, "metal"].values

comparison = df_test.groupby("n_atoms").agg({
    "Actual_DFT": "mean",
    "Predicted_ML": "mean",
    "Error": "mean"
}).round(5)
print(comparison)

# =========================================
print("\n===  Δ²E ===")
magic = df.groupby("n_atoms")["binding_energy_per_atom"].mean().reset_index().sort_values("n_atoms")

magic["Delta2E"] = magic["binding_energy_per_atom"].shift(-1) + magic["binding_energy_per_atom"].shift(1) - 2 * magic["binding_energy_per_atom"]

magic_numbers = magic[magic["Delta2E"] > 0.025]["n_atoms"].astype(int).tolist()

print("(Δ²E > 0.025):", magic_numbers)

original_magic = [6, 8, 12, 14, 18, 20, 32, 34, 38, 49, 55]

all_n = sorted(set(original_magic + magic_numbers))
table_data = []

for n in all_n:
    row = {"N": n}
    if n in magic["n_atoms"].values:
        eb = magic[magic["n_atoms"] == n]["binding_energy_per_atom"].values[0]
        delta2 = magic[magic["n_atoms"] == n]["Delta2E"].values[0]
        row["Avg_BE (eV/atom)"] = round(eb, 4)
        row["Δ²E (eV)"] = round(delta2, 4) if not np.isnan(delta2) else "N/A"
    else:
        row["Avg_BE (eV/atom)"] = "N/A"
        row["Δ²E (eV)"] = "N/A"
    
    if n in original_magic and n in magic_numbers:
        row["Type"] = "Both"
    elif n in original_magic:
        row["Type"] = "Original (Literature)"
    else:
        row["Type"] = "Discovered from QCD"
    table_data.append(row)

magic_table = pd.DataFrame(table_data)
print("\n=====")
print(magic_table.to_string(index=False))

# ===========================================
doc = Document()
doc.add_heading('Machine Learning Prediction of Coinage Metal Nanocluster Stability (Cu–Ag–Au) - N ≤ 55', 0)
doc.add_paragraph(f"Best model: **{best_model_name}** | MAE = {results_df.iloc[0]['MAE']:.5f} eV/atom | R² = {results_df.iloc[0]['R2']:.4f}")

# Table 1: Model Performance
doc.add_heading('Table 1: Performance of the 7 Models', level=2)
table = doc.add_table(rows=1, cols=4)
hdr = table.rows[0].cells
hdr[0].text = "Model"
hdr[1].text = "MAE (eV/atom)"
hdr[2].text = "R²"
hdr[3].text = "Time (s)"
for _, row in results_df.iterrows():
    r = table.add_row().cells
    r[0].text = row["Model"]
    r[1].text = f"{row['MAE']:.5f}"
    r[2].text = f"{row['R2']:.4f}"
    r[3].text = f"{row['Training_Time_sec']:.2f}"

# Table 2: ML vs DFT Comparison
doc.add_heading('Table 2: ML Predictions vs DFT Values (by cluster size)', level=2)
table2 = doc.add_table(rows=1, cols=4)
hdr = table2.rows[0].cells
hdr[0].text = "N"
hdr[1].text = "Avg DFT (eV/atom)"
hdr[2].text = "Avg ML (eV/atom)"
hdr[3].text = "MAE (eV/atom)"
for idx, row in comparison.iterrows():
    r = table2.add_row().cells
    r[0].text = str(idx)
    r[1].text = f"{row['Actual_DFT']:.4f}"
    r[2].text = f"{row['Predicted_ML']:.4f}"
    r[3].text = f"{row['Error']:.4f}"

# Table 3: Magic Numbers Comparison
doc.add_heading('Table 3: Magic Numbers Comparison (Original vs Discovered)', level=2)
table3 = doc.add_table(rows=1, cols=4)
hdr = table3.rows[0].cells
hdr[0].text = "N"
hdr[1].text = "Avg_BE (eV/atom)"
hdr[2].text = "Δ²E (eV)"
hdr[3].text = "Type"
for _, row in magic_table.iterrows():
    r = table3.add_row().cells
    r[0].text = str(row["N"])
    r[1].text = str(row["Avg_BE (eV/atom)"])
    r[2].text = str(row["Δ²E (eV)"])
    r[3].text = row["Type"]

doc.save(f"{ROOT}/Full_Paper_Report_N55_{datetime.now().strftime('%Y%m%d')}.docx")
print(f"\n Complete")

# ===============================================
def save_fig(fig, num, title):
    path = f"{ROOT}/Figures/Fig_{num:02d}_{title.replace(' ', '_')}.png"
    fig.savefig(path, dpi=600, bbox_inches='tight')
    plt.close(fig)
    print(f"Figure {num:02d} saved: {title}")

# Figure 1. Parity Plot
fig, ax = plt.subplots()
ax.scatter(y_test, best_model.predict(X_test_s), alpha=0.7, s=15)
ax.plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()], 'r--', lw=2)
ax.set_xlabel("Actual BE per atom (eV)")
ax.set_ylabel("Predicted BE per atom (eV)")
ax.set_title("Parity Plot")
save_fig(fig, 1, "Parity Plot")

# Figure 2. Stability vs Size
fig, ax = plt.subplots()
colors = {"Cu":"#1B5E20", "Ag":"#0D47A1", "Au":"#E65100"}
for m in ["Cu","Ag","Au"]:
    sub = df[df["metal"]==m]
    ax.scatter(sub["n_atoms"], sub["binding_energy_per_atom"], label=m, color=colors[m], s=15, alpha=0.7)
ax.set_xlabel("Cluster Size (n)")
ax.set_ylabel("Binding Energy per Atom (eV)")
ax.set_title("Stability vs Cluster Size")
ax.legend()
save_fig(fig, 2, "Stability vs Size")

# Figure 3. Feature Importance
imp = pd.Series(best_model.feature_importances_, index=feature_cols).sort_values(ascending=False)
fig, ax = plt.subplots()
imp.head(12).plot(kind='bar', ax=ax, color='#16a085')
ax.set_title("Feature Importance")
save_fig(fig, 3, "Feature Importance")

# Figure 4. SHAP
explainer = shap.TreeExplainer(best_model)
shap_values = explainer.shap_values(X_train_s[:500])
fig = plt.figure(figsize=(10,8))
shap.summary_plot(shap_values, X_train[:500], feature_names=feature_cols, show=False)
plt.title("SHAP Analysis")
save_fig(fig, 4, "SHAP Analysis")

# Figure 5. Learning Curve
train_sizes, train_scores, val_scores = learning_curve(best_model, X_train_s, y_train, cv=5,
                                                       train_sizes=np.linspace(0.1,1.0,6), scoring='neg_mean_absolute_error', n_jobs=-1)
fig, ax = plt.subplots()
ax.plot(train_sizes, -train_scores.mean(axis=1), 'o-', label='Training')
ax.plot(train_sizes, -val_scores.mean(axis=1), 'o-', label='Validation')
ax.set_xlabel("Training Size")
ax.set_ylabel("MAE (eV/atom)")
ax.set_title("Learning Curve")
ax.legend()
save_fig(fig, 5, "Learning Curve")

# Figure 6. Magic Numbers
magic = df.groupby("n_atoms")["binding_energy_per_atom"].mean().reset_index()
peaks, _ = find_peaks(magic["binding_energy_per_atom"], distance=2)
fig, ax = plt.subplots()
ax.plot(magic["n_atoms"], magic["binding_energy_per_atom"], marker='o')
for p in peaks:
    ax.scatter(magic.iloc[p]["n_atoms"], magic.iloc[p]["binding_energy_per_atom"], color='red', s=120)
ax.set_xlabel("Cluster Size")
ax.set_ylabel("Average BE per Atom")
ax.set_title("Magic Numbers")
save_fig(fig, 6, "Magic Numbers")

# Figure 7. Error by Metal
df_test = X_test.copy()
df_test["Actual"] = y_test.values
df_test["Predicted"] = best_model.predict(X_test_s)
df_test["Error"] = np.abs(df_test["Actual"] - df_test["Predicted"])
df_test["metal"] = df.loc[X_test.index, "metal"].values
fig, ax = plt.subplots()
sns.boxplot(x="metal", y="Error", data=df_test, palette=colors, ax=ax)
ax.set_ylabel("Absolute Error (eV/atom)")
ax.set_title("Error Distribution by Metal")
save_fig(fig, 7, "Error by Metal")

# Figure 8. Correlation Heatmap
fig, ax = plt.subplots(figsize=(8,6))
sns.heatmap(X.corr(), annot=True, fmt=".2f", cmap='coolwarm', ax=ax)
ax.set_title("Feature Correlation Heatmap")
save_fig(fig, 8, "Correlation Heatmap")

# Figure 9. PCA
from sklearn.decomposition import PCA
pca = PCA(n_components=2).fit_transform(X)
fig, ax = plt.subplots()
for m in ["Cu","Ag","Au"]:
    idx = df["metal"] == m
    ax.scatter(pca[idx,0], pca[idx,1], label=m, alpha=0.7, s=15, color=colors[m])
ax.set_xlabel("PC1")
ax.set_ylabel("PC2")
ax.set_title("PCA Projection")
ax.legend()
save_fig(fig, 9, "PCA Projection")

# Figure 10. t-SNE
from sklearn.manifold import TSNE
tsne = TSNE(n_components=2, random_state=42).fit_transform(X)
fig, ax = plt.subplots()
for m in ["Cu","Ag","Au"]:
    idx = df["metal"] == m
    ax.scatter(tsne[idx,0], tsne[idx,1], label=m, alpha=0.7, s=15, color=colors[m])
ax.set_xlabel("t-SNE 1")
ax.set_ylabel("t-SNE 2")
ax.set_title("t-SNE Manifold")
ax.legend()
save_fig(fig, 10, "t-SNE Projection")

# Figure 11: Residual Distribution
residuals = y_test - best_model.predict(X_test_s)
fig, ax = plt.subplots(figsize=(5.8, 4.3))
sns.histplot(residuals, kde=True, color='#9C27B0', bins=30, ax=ax)
ax.set_xlabel("Residual (Actual - Predicted) (eV/atom)")
ax.set_ylabel("Frequency")
ax.set_title("Residual Distribution")
save_fig(fig, 11, "Residual Distribution")

# Figure 12: Error vs Cluster Size
df_test = X_test.copy()
df_test["Actual"] = y_test.values
df_test["Predicted"] = best_model.predict(X_test_s)
df_test["Error"] = np.abs(df_test["Actual"] - df_test["Predicted"])
df_test["n_atoms"] = df.loc[X_test.index, "n_atoms"].values
df_test["metal"] = df.loc[X_test.index, "metal"].values

fig, ax = plt.subplots(figsize=(5.8, 4.3))
sns.scatterplot(data=df_test, x="n_atoms", y="Error", hue="metal", 
                palette=colors, s=25, alpha=0.7, ax=ax)
ax.set_xlabel("Cluster Size (n)")
ax.set_ylabel("Absolute Error (eV/atom)")
ax.set_title("Prediction Error vs Cluster Size")
save_fig(fig, 12, "Error vs Cluster Size")

# Figure 13: HOMO-LUMO Gap Distribution
fig, ax = plt.subplots(figsize=(5.8, 4.3))
sns.violinplot(data=df, x="metal", y="homo_lumo_gap", palette=colors, ax=ax)
ax.set_ylabel("HOMO-LUMO Gap (eV)")
ax.set_title("HOMO-LUMO Gap Distribution by Metal")
save_fig(fig, 13, "HOMO-LUMO Gap Distribution")

# Figure 14: Radius of Gyration vs Cluster Size
fig, ax = plt.subplots(figsize=(5.8, 4.3))
sns.scatterplot(data=df, x="n_atoms", y="radius_gyration", hue="metal", 
                palette=colors, s=20, alpha=0.7, ax=ax)
ax.set_xlabel("Cluster Size (n)")
ax.set_ylabel("Radius of Gyration (Å)")
ax.set_title("Radius of Gyration vs Cluster Size")
save_fig(fig, 14, "Radius of Gyration vs Size")

# Figure 15: Compactness Index
fig, ax = plt.subplots(figsize=(5.8, 4.3))
sns.boxplot(data=df, x="metal", y="compactness", palette=colors, ax=ax)
ax.set_ylabel("Compactness Index")
ax.set_title("Compactness Index by Metal")
save_fig(fig, 15, "Compactness Index")

# Figure 16: Asphericity Distribution
fig, ax = plt.subplots(figsize=(5.8, 4.3))
sns.kdeplot(data=df, x="asphericity", hue="metal", fill=True, palette=colors, ax=ax)
ax.set_xlabel("Asphericity")
ax.set_title("Asphericity Distribution by Metal")
save_fig(fig, 16, "Asphericity Distribution")

# Figure 17: Radius of Gyration vs Cluster Size
fig, ax = plt.subplots(figsize=(5.8, 4.3))
for m in ["Cu", "Ag", "Au"]:
    sub = df[df["metal"] == m]
    ax.scatter(sub["n_atoms"], sub["radius_gyration"], label=m, color=colors[m], s=15, alpha=0.6)
ax.set_xlabel("Cluster Size (n)")
ax.set_ylabel("Radius of Gyration (Å)")
ax.set_title("Radius of Gyration vs Cluster Size")
ax.legend()
save_fig(fig, 17, "Radius of Gyration vs Size")

# Figure 18: Compactness Distribution
fig, ax = plt.subplots(figsize=(5.8, 4.3))
sns.violinplot(data=df, x="metal", y="compactness", palette=colors, ax=ax)
ax.set_ylabel("Compactness")
ax.set_title("Compactness Distribution")
save_fig(fig, 18, "Compactness Distribution")

# Figure 19: Energy Density Distribution
fig, ax = plt.subplots(figsize=(5.8, 4.3))
sns.kdeplot(data=df, x="binding_energy_per_atom", hue="metal", fill=True, palette=colors, ax=ax)
ax.set_xlabel("Binding Energy per Atom (eV)")
ax.set_title("Energy Density Distribution")
save_fig(fig, 19, "Energy Density Distribution")

# Figure 20: Asphericity (Shape Deviation)
fig, ax = plt.subplots(figsize=(5.8, 4.3))
sns.stripplot(data=df, x="metal", y="asphericity", palette=colors, alpha=0.6, ax=ax)
ax.set_ylabel("Asphericity")
ax.set_title("Shape Deviation (Asphericity)")
save_fig(fig, 20, "Asphericity Distribution")

# Figure 21: Bounding Box Growth
fig, ax = plt.subplots(figsize=(5.8, 4.3))
sns.lineplot(data=df, x="n_atoms", y="bbox_x", hue="metal", palette=colors, ax=ax)
ax.set_xlabel("Cluster Size (n)")
ax.set_ylabel("Bounding Box X (Å)")
ax.set_title("Bounding Box Growth")
save_fig(fig, 21, "Bounding Box Growth")

# Figure 22: Lowess Non-parametric Fit
y_pred = best_model.predict(X_test_s)
from statsmodels.nonparametric.smoothers_lowess import lowess
lowess_fit = lowess(y_pred, y_test, frac=0.3)
fig, ax = plt.subplots(figsize=(5.8, 4.3))
ax.scatter(y_test, y_pred, alpha=0.5, s=15)
ax.plot(lowess_fit[:,0], lowess_fit[:,1], color='red', lw=2.5)
ax.set_xlabel("Actual BE per atom (eV)")
ax.set_ylabel("Predicted BE per atom (eV)")
ax.set_title("Lowess Non-parametric Fit")
save_fig(fig, 22, "Lowess Fit")

# Figure 23: 3D Stability Landscape
fig = plt.figure(figsize=(8, 6))
ax = fig.add_subplot(111, projection='3d')
sc = ax.scatter(df["n_atoms"], df["radius_gyration"], df["binding_energy_per_atom"],
                c=df["metal_Z"], cmap="viridis", s=20, alpha=0.8)
ax.set_xlabel("Cluster Size")
ax.set_ylabel("Radius of Gyration (Å)")
ax.set_zlabel("Binding Energy per Atom (eV)")
plt.colorbar(sc, label="Atomic Number (Z)")
ax.set_title("3D Stability Landscape")
save_fig(fig, 23, "3D Stability Landscape")

# Figure 24: Binding Energy by Metal
fig, ax = plt.subplots(figsize=(5.8, 4.3))
sns.boxplot(data=df, x="metal", y="binding_energy_per_atom", palette=colors, ax=ax)
ax.set_ylabel("Binding Energy per Atom (eV)")
ax.set_title("Binding Energy Distribution by Metal")
save_fig(fig, 24, "BE by Metal")

# Figure 25: Global Feature Importance
fig, ax = plt.subplots(figsize=(8, 7))
imp = pd.Series(best_model.feature_importances_, index=feature_cols).sort_values(ascending=True)
imp.plot(kind='barh', ax=ax, color='#16a085')
ax.set_title("Global Feature Importance")
save_fig(fig, 25, "Global Feature Importance")

# Figure 26: Full Correlation Matrix
fig, ax = plt.subplots(figsize=(9, 7))
sns.heatmap(X.corr(), annot=True, fmt=".2f", cmap='RdBu_r', center=0, ax=ax, annot_kws={'size': 8})
ax.set_title("Full Feature Correlation Matrix")
save_fig(fig, 26, "Full Correlation Matrix")

print("\n Compelet")

# ====================== (Zip) ======================
print("\n" + "="*90)
zip_name = f"{ROOT}_Final_{datetime.now().strftime('%Y%m%d_%H%M')}"
zip_path = shutil.make_archive(zip_name, 'zip', ROOT)

size_mb = os.path.getsize(zip_path) / (1024 * 1024)
print(f"Comeplet: {zip_path} ({size_mb:.1f} MB)")

print("\n!")


